# -*- coding: utf-8 -*-
"""将测试截图与修复说明打包成 .doc 报告"""
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT = Path(r"D:\虚拟机\期末考核\NetForum")
SHOT_DIR = ROOT / "screenshots"
OUT = ROOT / "NetForum_修复报告.docx"

doc = Document()

# 全局样式
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


def add_h1(text):
    p = doc.add_heading(level=1)
    p.add_run(text)
    return p


def add_h2(text):
    p = doc.add_heading(level=2)
    p.add_run(text)
    return p


def add_p(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    return p


def add_code(path, content):
    """插入一段代码片段"""
    add_p(f"文件: {path}", bold=True)
    p = doc.add_paragraph()
    r = p.add_run(content)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    pPr = p._p.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:fill'): 'F4F4F4'})
    pPr.append(shd)


def add_img(path, caption):
    if Path(path).exists():
        doc.add_picture(str(path), width=Inches(6.0))
        # 居中
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph()
        r = p.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_p(f"[图片缺失] {path}")


# ============== 封面 ==============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("NetForum 网上论坛系统")
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0xB1, 0x1A, 0x1A)
r.bold = True

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Bug 修复与测试报告")
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
r.bold = True

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("修复日期：2026-06-15\n")
meta.add_run("测试环境：JDK 21 + Maven 3.9.6 + Tomcat 9.0.117 + MySQL 9.5\n")
meta.add_run("测试账号：admin / admin123（管理员）\n")
meta.add_run("浏览器：Chromium（Playwright 自动化）")

doc.add_page_break()

# ============== 一、问题概览 ==============
add_h1("一、问题概览")
add_p("本次共修复 3 个功能 Bug，并附带解决 1 个隐藏的次生 Bug。")

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = "编号"
hdr[1].text = "Bug 描述"
hdr[2].text = "影响范围"
hdr[3].text = "严重度"
rows = [
    ("Bug #1", "管理员编辑非本人帖子时，直接被重定向到首页", "管理员日常运营", "高"),
    ("Bug #2", "管理员删除带回复的帖子时，失败并返回一段 HTML/JSON 源码", "管理员日常运营", "高"),
    ("Bug #3", "底部导航栏「使用指南/社区公约/联系我们」使用 javascript:void(0) 占位", "所有用户", "中"),
    ("Bug #4（次生）", "handleUpdate 没有真正调用 postService.updatePost，编辑后标题未入库", "所有用户编辑", "高"),
]
for r in rows:
    cells = table.add_row().cells
    for i, v in enumerate(r):
        cells[i].text = v

doc.add_page_break()

# ============== 二、Bug #1 修复 ==============
add_h1("二、Bug #1 — 管理员编辑帖子跳首页")
add_h2("根因分析")
add_p("PostServlet.doGet() 的 /edit 分支只允许帖子作者编辑，未放行管理员：")
add_code(
    "src/main/java/com/netforum/servlet/PostServlet.java (第 53 行，修复前)",
    "// ❌ 仅作者可以编辑，管理员会被踢到首页\n"
    "if (loginUser != null && loginUser.getId().equals(post.getUserId())) {\n"
    "    req.setAttribute(\"post\", post);\n"
    "    ...\n"
    "}\n"
    "resp.sendRedirect(req.getContextPath() + \"/\");  // 管理员会命中这一行"
)

add_h2("修复方案")
add_p("在原条件后追加 isAdmin() 放行：")
add_code(
    "src/main/java/com/netforum/servlet/PostServlet.java (第 53 行，修复后)",
    "// ✅ 作者或管理员都可以编辑\n"
    "if (loginUser != null &&\n"
    "    (loginUser.getId().equals(post.getUserId()) || loginUser.isAdmin())) {\n"
    "    req.setAttribute(\"post\", post);\n"
    "    ...\n"
    "}\n"
    "resp.sendRedirect(req.getContextPath() + \"/\");"
)
add_p("顺便同步修复了同文件 handleUpdate 中 if 条件优先级不清的问题（加括号）：")
add_code(
    "src/main/java/com/netforum/servlet/PostServlet.java (handleUpdate 修复后)",
    "if (post == null) { resp.sendRedirect(...); return; }\n"
    "// 权限校验：作者或管理员\n"
    "if (!post.getUserId().equals(loginUser.getId()) && !loginUser.isAdmin()) {\n"
    "    resp.sendRedirect(...); return;\n"
    "}\n"
    "// 真正调用持久化（修复 Bug #4）\n"
    "boolean ok = postService.updatePost(post.getId(), title, content);"
)

add_h2("测试截图")
add_p("测试场景：admin 登录后点击其他用户（user_id=5）的帖子 id=4 的「编辑」按钮，截图如下：")
add_img(SHOT_DIR / "07_admin_edit_post_not_own.png",
        "图 1：admin 进入帖子 id=4（作者非 admin）的编辑页，未再跳到首页")
add_img(SHOT_DIR / "08_admin_edit_post_result.png",
        "图 2：admin 修改标题后提交，详情页正确显示新标题「管理员测试编辑-原标题」")

doc.add_page_break()

# ============== 三、Bug #2 修复 ==============
add_h1("三、Bug #2 — 删除帖子失败显示 HTML 源码")
add_h2("根因分析")
add_p("经排查，根因有两层：")
add_p("1. 数据库外键约束：t_reply.post_id / t_attachment.post_id 都外键引用 t_post.id，"
      "因此当一个帖子拥有回复时，直接 DELETE FROM t_post 会被 MySQL 拒绝；"
      "PostService.deletePost() 之前只级联删了 t_like，并未删 t_reply/t_attachment，"
      "导致 ps.executeUpdate() 返回 0，被判定为「删除失败」。", bold=True)
add_p("2. 失败响应处理不当：PostServlet.handleDelete 在删除失败时调用 "
      "resp.getWriter().write(\"{...}\") 输出 JSON，"
      "但调用方是浏览器表单 POST 提交，期待重定向到下一个 HTML 页面，"
      "于是浏览器把 JSON 字符串原样展示出来。", bold=True)
add_code(
    "src/main/java/com/netforum/service/PostService.java (修复前 deletePost)",
    "public boolean deletePost(Integer postId) {\n"
    "    Post post = postDAO.findById(postId);\n"
    "    if (post != null) {\n"
    "        likeDAO.removeLikesByPostId(postId);   // ❌ 只删了点赞\n"
    "        boolean deleted = postDAO.delete(postId);  // 触发外键违反\n"
    "        ...\n"
    "    }\n"
    "}"
)

add_h2("修复方案")
add_p("(a) 在 ReplyDAO / AttachmentDAO 新增 deleteByPostId：")
add_code(
    "src/main/java/com/netforum/dao/ReplyDAO.java (新增)",
    "public boolean deleteByPostId(Integer postId) {\n"
    "    String sql = \"DELETE FROM t_reply WHERE post_id=?\";\n"
    "    ...\n"
    "}"
)
add_p("(b) PostService.deletePost 改为四级级联删除：")
add_code(
    "src/main/java/com/netforum/service/PostService.java (修复后 deletePost)",
    "public boolean deletePost(Integer postId) {\n"
    "    Post post = postDAO.findById(postId);\n"
    "    if (post != null) {\n"
    "        replyDAO.deleteByPostId(postId);          // 1. 删回复\n"
    "        attachmentDAO.deleteByPostId(postId);     // 2. 删附件\n"
    "        likeDAO.removeLikesByPostId(postId);       // 3. 删点赞\n"
    "        boolean deleted = postDAO.delete(postId);  // 4. 删帖子\n"
    "        if (deleted) boardDAO.updatePostCount(post.getBoardId(), -1);\n"
    "        return deleted;\n"
    "    }\n"
    "    return false;\n"
    "}"
)
add_p("(c) PostServlet.handleDelete 在失败时改为重定向到详情页 + flash 提示，"
      "让浏览器看到的是正常 HTML 页面：")
add_code(
    "src/main/java/com/netforum/servlet/PostServlet.java (handleDelete 修复后)",
    "if (success) {\n"
    "    resp.sendRedirect(req.getContextPath() + \"/board/detail?id=\" + post.getBoardId());\n"
    "} else {\n"
    "    req.getSession().setAttribute(\"flashError\", \"删除帖子失败，请稍后重试\");\n"
    "    resp.sendRedirect(req.getContextPath() + \"/post/detail?id=\" + post.getId());\n"
    "}"
)
add_p("(d) post.jsp / edit_post.jsp 顶部读取并显示 flashError，配以 alert CSS 样式。")

add_h2("测试截图")
add_img(SHOT_DIR / "09_delete_success.png",
        "图 3：admin 删除带 6 条回复的帖子 id=31 后，正确跳转到「生活休闲」板块页（无 JSON 源码）")
add_img(SHOT_DIR / "10_post_deleted.png",
        "图 4：再次访问被删帖子 id=31，重定向到首页（确认数据库记录已删除）")

doc.add_page_break()

# ============== 四、Bug #3 修复 ==============
add_h1("四、Bug #3 — 底部导航栏超链接无法点击")
add_h2("根因分析")
add_p("footer.jsp 中「使用指南」「社区公约」「联系我们」三个链接使用了 javascript:void(0) 占位：")
add_code(
    "src/main/webapp/pages/common/footer.jsp (修复前)",
    "<li><a href=\"javascript:void(0)\">使用指南</a></li>\n"
    "<li><a href=\"javascript:void(0)\">社区公约</a></li>\n"
    "<li><a href=\"javascript:void(0)\">联系我们</a></li>"
)
add_p("点击后没有任何响应，用户体验受损。")

add_h2("修复方案")
add_p("新增一个轻量的 PageServlet 做路由 (/page/*)，并创建三个真实页面：")
add_code(
    "src/main/java/com/netforum/servlet/PageServlet.java (新增)",
    "@WebServlet(\"/page/*\")\n"
    "public class PageServlet extends HttpServlet {\n"
    "    protected void doGet(req, resp) {\n"
    "        switch (req.getPathInfo()) {\n"
    "            case \"/help\":       jspPath = \"/pages/static/help.jsp\";       break;\n"
    "            case \"/agreement\":  jspPath = \"/pages/static/agreement.jsp\";  break;\n"
    "            case \"/contact\":    jspPath = \"/pages/static/contact.jsp\";    break;\n"
    "        }\n"
    "        req.getRequestDispatcher(jspPath).forward(req, resp);\n"
    "    }\n"
    "}"
)
add_p("footer.jsp 改为指向真实路径：")
add_code(
    "src/main/webapp/pages/common/footer.jsp (修复后)",
    "<li><a href=\"${pageContext.request.contextPath}/page/help\">使用指南</a></li>\n"
    "<li><a href=\"${pageContext.request.contextPath}/page/agreement\">社区公约</a></li>\n"
    "<li><a href=\"${pageContext.request.contextPath}/page/contact\">联系我们</a></li>"
)
add_p("新建文件：")
add_p("• src/main/webapp/pages/static/help.jsp       （使用指南）")
add_p("• src/main/webapp/pages/static/agreement.jsp  （社区公约）")
add_p("• src/main/webapp/pages/static/contact.jsp    （联系我们）")
add_p("css/style.css 末尾补充 .static-page 与 .alert--error / .alert--success 样式。")

add_h2("测试截图")
add_img(SHOT_DIR / "01_home_with_new_footer.png",
        "图 5：首页底部「关于」列的三个链接已替换为真实可点击链接")
add_img(SHOT_DIR / "02_footer_help.png",
        "图 6：点击「使用指南」 → /page/help 正常渲染使用指南页面")
add_img(SHOT_DIR / "03_footer_agreement.png",
        "图 7：点击「社区公约」 → /page/agreement 正常渲染社区公约页面")
add_img(SHOT_DIR / "04_footer_contact.png",
        "图 8：点击「联系我们」 → /page/contact 正常渲染联系我们页面")

doc.add_page_break()

# ============== 五、编译部署 ==============
add_h1("五、编译与部署")
add_p("Maven 编译输出：")
add_code(
    "mvn clean package -DskipTests",
    "[INFO] --- compiler:3.11.0:compile (default-compile) @ netforum ---\n"
    "[INFO] Compiling 30 source files with javac [debug target 11] to target\\classes\n"
    "[INFO] ---\n"
    "[INFO] --- war:3.4.0:war (default-war) @ netforum ---\n"
    "[INFO] Building war: D:\\...\\target\\netforum.war\n"
    "[INFO] BUILD SUCCESS\n"
    "[INFO] Total time:  3.457 s"
)
add_p("部署到 Tomcat 9.0.117，Web 应用程序 netforum.war 启动正常。")

# ============== 六、回归验证 ==============
add_h1("六、自动化回归测试")
add_p("使用 Playwright 1.60 + Chromium headless 浏览器对三个修复做了端到端测试，"
      "测试脚本保存为 test_fixes.py。")
add_p("测试流程：")
add_p("1. 验证底部导航栏三个链接全部可点击并正常渲染", bold=False)
add_p("2. 用 admin 账号登录，验证登录后正常显示用户名", bold=False)
add_p("3. admin 进入帖子 id=4（作者非 admin）的编辑页，断言 URL 包含 /post/edit "
      "且页面显示编辑表单（不再跳首页）", bold=False)
add_p("4. 修改标题并提交，断言详情页内容包含新标题", bold=False)
add_p("5. admin 删除带 6 条回复的帖子 id=31，断言响应为 302 重定向、跳转板块页、"
      "且数据库中该帖子已被删除", bold=False)

add_h2("关键测试结果")
add_code(
    "playwright run output (节选)",
    "[测试1] 管理员编辑非本人帖子\n"
    "  编辑帖子 id=4 后 url: http://localhost:8080/netforum/post/edit?id=4\n"
    "  提交编辑后 url: http://localhost:8080/netforum/post/detail?id=4\n"
    "  提交编辑后包含新标题: True     ← 修复前是 False\n"
    "\n"
    "[测试2] 管理员删除带回复的帖子\n"
    "  详情页 url: http://localhost:8080/netforum/post/detail?id=31\n"
    "  拦截到 /post/delete 响应: status=302 url=.../post/delete\n"
    "  删除后 url: http://localhost:8080/netforum/board/detail?id=2\n"
    "  ✓ 删除成功跳转到板块页         ← 修复前是显示 JSON 源码\n"
    "  ✓ 帖子已被删除，重定向到: http://localhost:8080/netforum/\n"
)

# ============== 七、修改清单 ==============
add_h1("七、本次修改文件清单")
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = "文件"
hdr[1].text = "类型"
hdr[2].text = "说明"
rows = [
    ("src/main/java/com/netforum/servlet/PostServlet.java", "修改",
     "放行管理员编辑、修正 handleUpdate 真正调用持久化、handleDelete 改重定向"),
    ("src/main/java/com/netforum/servlet/PageServlet.java", "新增",
     "/page/* 路由，转发到三个静态页 JSP"),
    ("src/main/java/com/netforum/dao/PostDAO.java", "修改", "新增 update(id,title,content)"),
    ("src/main/java/com/netforum/dao/ReplyDAO.java", "修改", "新增 deleteByPostId(postId)"),
    ("src/main/java/com/netforum/dao/AttachmentDAO.java", "修改", "新增 deleteByPostId(postId)"),
    ("src/main/java/com/netforum/service/PostService.java", "修改",
     "deletePost 四级级联删，新增 updatePost"),
    ("src/main/webapp/pages/common/footer.jsp", "修改",
     "javascript:void(0) → /page/help 等真实链接"),
    ("src/main/webapp/pages/post/post.jsp", "修改",
     "顶部读取并显示 flashError 错误提示"),
    ("src/main/webapp/pages/post/edit_post.jsp", "修改",
     "顶部读取并显示 flashError 错误提示"),
    ("src/main/webapp/pages/static/help.jsp", "新增", "使用指南页"),
    ("src/main/webapp/pages/static/agreement.jsp", "新增", "社区公约页"),
    ("src/main/webapp/pages/static/contact.jsp", "新增", "联系我们页"),
    ("src/main/webapp/css/style.css", "修改", "新增 .static-page、.alert--error 样式"),
]
for r in rows:
    cells = table.add_row().cells
    for i, v in enumerate(r):
        cells[i].text = v

doc.add_page_break()

# ============== 八、总结 ==============
add_h1("八、总结")
add_p("✅ Bug #1 修复成功：管理员可正常编辑非本人帖子。")
add_p("✅ Bug #2 修复成功：删除带回复的帖子不再失败，浏览器也不会再看到 JSON 源码。")
add_p("✅ Bug #3 修复成功：底部导航栏三个链接全部可点击并正常渲染。")
add_p("✅ 顺带修复 Bug #4：handleUpdate 现在会真正调用持久化，编辑后标题/内容立即生效。")
add_p("")
add_p("⚠ 关注：", bold=True)
add_p("• 回复和附件的级联删除目前是 Application 层手动级联，"
      "若新增与 t_post 关联的子表，需同步在 PostService.deletePost 中添加对应删除逻辑，"
      "或考虑在数据库层加 ON DELETE CASCADE 外键。", bold=False)
add_p("")
add_p("💡 改进建议：", bold=True)
add_p("• 静态页面（使用指南/社区公约/联系我们）目前以 JSP 形式硬编码，未来可改造为 "
      "管理员可在后台编辑的 CMS 模块。", bold=False)
add_p("• 编辑帖子的前后端可加一个 Markdown 实时预览面板，提升体验。", bold=False)

doc.save(str(OUT))
print(f"已生成报告: {OUT}")
print(f"文件大小: {OUT.stat().st_size / 1024:.1f} KB")
